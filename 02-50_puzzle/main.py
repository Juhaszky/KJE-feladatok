from docx import Document
import sys
from io import StringIO
import math
import subprocess
from docx.enum.text import WD_ALIGN_PARAGRAPH


def main():
    document = readDocument()
    answers = process_paragraphs(document.paragraphs)
    print('Docx excercises solved!')
    fill_solution_lines(answers)
    print('With_solution docx created, filled with answers!')


def readDocument():
    return Document("50_puzzle.docx")

def process_paragraphs(paragraphs):
    answers = []
    python_code_lines = []
    for paragraph in paragraphs:
        if paragraph.text.startswith("What is the output of this code?"):
            python_code_lines = []
        elif paragraph.text.startswith("The correct solution:"):
            code = "\n".join(python_code_lines)
            result_value = execute_with_timeout(code, timeout=2)
            if (result_value != ""):
                answers.append(result_value)
            else:
                answers.append("Syntax error!")
        else:
            paragraph_first_line_indent_value, paragraph_left_indent_value = get_paragraph_indent_values(paragraph)
            if paragraph_first_line_indent_value > 0:
                if paragraph_left_indent_value > 0:
                    p_text = get_tabulators(paragraph_first_line_indent_value + paragraph_left_indent_value) + paragraph.text
                    python_code_lines.append(p_text)
                else:
                    p_text = get_tabulators(paragraph_first_line_indent_value) + paragraph.text
                    python_code_lines.append(p_text)
            else: 
                python_code_lines.append(paragraph.text)      
    return answers


def fill_solution_lines(answers):
    document = Document("50_puzzle.docx")
    task_counter = 0
    for paragraph in document.paragraphs:
        if ("The correct solution:" in paragraph.text):
            if (answers[task_counter] is not None):
                paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
                paragraph.text = paragraph.text.strip() + "\n" + answers[task_counter].strip()
                task_counter += 1
    document.save('50_puzzle_with_solutions.docx')

def get_paragraph_indent_values(paragraph):
    paragraph_first_line_indent_value = 0
    paragraph_left_indent_value = 0
    
    if (paragraph.paragraph_format.first_line_indent is not None):
        paragraph_first_line_indent_value = paragraph.paragraph_format.first_line_indent.cm
    if (paragraph.paragraph_format.left_indent is not None):
        paragraph_left_indent_value = paragraph.paragraph_format.left_indent.cm

    return [paragraph_first_line_indent_value, paragraph_left_indent_value]

def get_tabulators(left_indent_amount):
    return "\t" * math.floor(left_indent_amount) 

def execute_with_timeout(code, timeout=1):
    try:
        result = subprocess.run(["python", "-c", code], capture_output=True, text=True, timeout=timeout)
        return result.stdout
    except subprocess.TimeoutExpired:
        return "Code execution timed out."
    except subprocess.CalledProcessError as e:
        return f"Syntax error in code: {e.stderr}"
    except Exception as e:
        return f"Error during code execution: {e}"

main()